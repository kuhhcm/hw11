import requests
import json
from docx import Document 



# url = 'https://jsonplaceholder.typicode.com/todos'
# response = requests.get(url)
# data = response.json()

# with open('todos.json', 'w') as file:
#     json.dump(data, file)

# with open('todos.json', 'r') as file:
#     data = json.load(file)

# for index, todo in enumerate(data):
#     file_name = f'todo_{index + 1}.json'
#     with open(file_name,'w') as todo_wile:
#         json.dump(todo, todo_wile)




# doc = Document()
# doc.add_paragraph("Hello Python")
# doc.save("hello_python.docx")

# doc = Document("hello_python.docx")
# bold_text = ""

# for paragraph in doc.paragraphs:
#     for run in paragraph.runs:
#         if run.bold:
#             bold_text += run.text

# print("Bold Text in Python String:", bold_text)
